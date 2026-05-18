"""
Document loader — ingests CSV, Excel, Word, PDF into a unified text format.
"""

import os
from dataclasses import dataclass, field
from pathlib import Path

import pandas as pd


@dataclass
class Document:
    text: str
    metadata: dict = field(default_factory=dict)


def load_csv(path: str) -> list[Document]:
    df = pd.read_csv(path)
    return _dataframe_to_docs(df, path)


def load_excel(path: str) -> list[Document]:
    xls = pd.ExcelFile(path)
    docs = []
    for sheet in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet)
        docs.extend(_dataframe_to_docs(df, path, sheet=sheet))
    return docs


def load_word(path: str) -> list[Document]:
    from docx import Document as DocxDocument

    doc = DocxDocument(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    full_text = "\n".join(paragraphs)
    if not full_text.strip():
        return []
    return [Document(text=full_text, metadata={"source": path, "type": "word"})]


def load_pdf(path: str) -> list[Document]:
    import pdfplumber

    docs = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            # Also try to extract tables
            tables = page.extract_tables() or []
            table_text = ""
            for table in tables:
                for row in table:
                    cleaned = [str(cell).strip() if cell else "" for cell in row]
                    table_text += " | ".join(cleaned) + "\n"

            combined = (text + "\n" + table_text).strip()
            if combined:
                docs.append(Document(
                    text=combined,
                    metadata={"source": path, "type": "pdf", "page": i + 1},
                ))
    return docs


def _dataframe_to_docs(
    df: pd.DataFrame, path: str, sheet: str | None = None
) -> list[Document]:
    """Convert each row of a DataFrame into a Document with column-value pairs."""
    docs = []
    columns = list(df.columns)
    for idx, row in df.iterrows():
        parts = []
        for col in columns:
            val = row[col]
            if pd.notna(val):
                parts.append(f"{col}: {val}")
        if parts:
            meta = {"source": path, "type": "tabular", "row": idx}
            if sheet:
                meta["sheet"] = sheet
            docs.append(Document(text="\n".join(parts), metadata=meta))
    return docs


LOADERS = {
    ".csv": load_csv,
    ".xlsx": load_excel,
    ".xls": load_excel,
    ".docx": load_word,
    ".doc": load_word,
    ".pdf": load_pdf,
}


def load_file(path: str) -> list[Document]:
    ext = Path(path).suffix.lower()
    loader = LOADERS.get(ext)
    if loader is None:
        raise ValueError(f"Unsupported file type: {ext}. Supported: {list(LOADERS)}")
    return loader(path)


def load_directory(directory: str) -> list[Document]:
    docs = []
    for root, _, files in os.walk(directory):
        for fname in sorted(files):
            ext = Path(fname).suffix.lower()
            if ext in LOADERS:
                full_path = os.path.join(root, fname)
                try:
                    docs.extend(load_file(full_path))
                except Exception as e:
                    print(f"Warning: failed to load {full_path}: {e}")
    return docs
